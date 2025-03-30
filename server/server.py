from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from flask_bcrypt import Bcrypt
from flask_jwt_extended import JWTManager, create_access_token
from pymongo import MongoClient
import os
from dotenv import load_dotenv
import tempfile
from werkzeug.utils import secure_filename
import PyPDF2
from io import BytesIO
from docx import Document as DocxDocument
import uuid
from datetime import timedelta

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)
app.config['JWT_SECRET_KEY'] = os.getenv('JWT_SECRET')
app.config['UPLOAD_FOLDER'] = tempfile.mkdtemp()
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB limit
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(hours=1)
jwt = JWTManager(app)
bcrypt = Bcrypt(app)

# Connect to MongoDB
client = MongoClient(os.getenv("MONGO_URI"))
db = client.get_database("InsightPaper")
users_collection = db["users"]
documents_collection = db["documents"]

# Allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_stream):
    try:
        pdf_reader = PyPDF2.PdfReader(file_stream)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        print(f"PDF extraction error: {str(e)}")
        return None

def extract_text_from_docx(file_stream):
    try:
        doc = DocxDocument(file_stream)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"DOCX extraction error: {str(e)}")
        return None

# Auth Routes
@app.route('/signup', methods=['POST'])
def signup():
    data = request.json
    username = data.get('username')
    email = data.get('email')
    password = data.get('password')

    if users_collection.find_one({"email": email}):
        return jsonify({"message": "User already exists"}), 400

    hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
    users_collection.insert_one({
        "username": username,
        "email": email,
        "password": hashed_password
    })
    
    return jsonify({"message": "User registered successfully"}), 201

@app.route('/login', methods=['POST'])
def login():
    data = request.json
    username = data.get('username')
    password = data.get('password')

    user = users_collection.find_one({"username": username})
    if not user or not bcrypt.check_password_hash(user['password'], password):
        return jsonify({"message": "Invalid credentials"}), 401

    token = create_access_token(identity=str(user['_id']))
    return jsonify({
        "message": "Login successful",
        "token": token,
        "user": {
            "id": str(user['_id']),
            "username": user['username'],
            "email": user['email']
        }
    })

# Document Routes
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    if not allowed_file(file.filename):
        return jsonify({"error": "Only PDF and DOCX files are allowed"}), 400
    
    try:
        # Save file with unique name
        file_ext = file.filename.rsplit('.', 1)[1].lower()
        unique_id = str(uuid.uuid4())
        filename = f"doc_{unique_id}.{file_ext}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # Store document metadata
        doc_data = {
            "original_name": file.filename,
            "stored_name": filename,
            "user_id": "temp_user",  # Replace with JWT user ID in production
            "upload_date": datetime.utcnow(),
            "file_type": file_ext,
            "size": os.path.getsize(filepath)
        }
        documents_collection.insert_one(doc_data)

        return jsonify({
            "message": "File uploaded successfully",
            "filename": file.filename,
            "preview_url": f"/preview/{filename}"
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/preview/<filename>', methods=['GET'])
def preview_document(filename):
    try:
        # Security check
        if not filename.startswith('doc_'):
            return jsonify({"error": "Invalid file"}), 400
        
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(filepath):
            return jsonify({"error": "File not found"}), 404
        
        # PDF preview
        if filename.lower().endswith('.pdf'):
            return send_from_directory(
                app.config['UPLOAD_FOLDER'],
                filename,
                as_attachment=False,
                mimetype='application/pdf'
            )
        
        # DOCX preview (extract text)
        elif filename.lower().endswith('.docx'):
            doc = DocxDocument(filepath)
            text = "\n".join([para.text for para in doc.paragraphs[:20]])  # First 20 paragraphs
            return jsonify({
                "type": "docx",
                "content": text,
                "filename": filename
            })
        
        return jsonify({"error": "Unsupported file type"}), 400

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/process', methods=['POST'])
def process_document():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    
    file = request.files['file']
    query = request.form.get('query', '').strip()
    
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    if not allowed_file(file.filename):
        return jsonify({"error": "Only PDF and DOCX files are allowed"}), 400
    
    try:
        # Extract text
        file_stream = BytesIO(file.read())
        file_ext = file.filename.rsplit('.', 1)[1].lower()
        
        if file_ext == 'pdf':
            text = extract_text_from_pdf(file_stream)
        elif file_ext == 'docx':
            text = extract_text_from_docx(file_stream)
        else:
            return jsonify({"error": "Unsupported file type"}), 400
        
        if not text:
            return jsonify({"error": "Could not extract text from document"}), 400
        
        # Mock processing - replace with actual ML/NLP processing
        if "summarize" in query.lower():
            response = "This is a mock summary of the document. The actual implementation would analyze the content and provide a concise summary."
        else:
            response = f"Mock response to: '{query}'. The document contains {len(text.split())} words and appears to be a {file_ext.upper()} file."
        
        return jsonify({
            "response": response,
            "word_count": len(text.split()),
            "file_type": file_ext
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=int(os.getenv("PORT", 5000)))